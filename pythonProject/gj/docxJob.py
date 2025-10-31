from fileinput import filename

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def todocx(docxName='监控一号',imgPath='./static/frontend_screenshot/screenshot_with_bboxes.png',filePath = './隐患详情分析.md',outPath='./static/frontend_screenshot/'):
    doc = Document()
    title = doc.add_heading('详情清单', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para1 = doc.add_paragraph(docxName+'：')
    para2 = doc.add_paragraph('监控标注图')
    doc.add_picture(imgPath, width=Inches(6))  # 宽度设为 2 英寸
    para3 = doc.add_paragraph('隐患详细分析')
    with open(filePath, 'r', encoding='utf-8') as f:
        text_content = f.read()
    para4 = doc.add_paragraph(text_content)

    doc.save(outPath+docxName+'文档报告.docx', )

def questionsText(filePath = '../static/job_screenshot/隐患详情分析.md'):
    with open(filePath,'r',encoding='utf-8') as f:
        text_content=f.read()
    textList=text_content.split('\n')
    questionList=[line[5::].replace(" ","") for line in textList if line.startswith('场景状况')]
    print(questionList)
    return questionList

def todocx2(docxName='监控一号',imgPath='../static/job_screenshot/screenshot.png',filePath = '../static/job_screenshot/隐患详情分析.md',outPath='../static/job_screenshot/'):
    doc = Document()
    title = doc.add_heading('问题清单', level=0)
    questionlist=questionsText(filePath)
    table = doc.add_table(rows=len(questionlist)+2, cols=2)
    table.cell(0,0).text='项目名称'
    table.cell(0,1).text='工地巡检项目'
    for i in range(len(questionlist)):
        table.cell(i+1,0).text='问题'+str(i+1)
        table.cell(i+1,1).text=questionlist[i]
    row=table.rows[len(questionlist)+2-1]
    merged_cell=row.cells[0].merge(row.cells[1])

    paragraph=merged_cell.add_paragraph()
    run=paragraph.add_run()

    run.add_picture(imgPath, width=Inches(6))

    doc.save(outPath+docxName+'问题报告.docx', )

